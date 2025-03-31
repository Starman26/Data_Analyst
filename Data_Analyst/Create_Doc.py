from docx import Document
from docx.shared import Inches, RGBColor, Pt
from datetime import datetime
import os

def create_docx(image_folder, save_folder="reportes"):
    os.makedirs(save_folder, exist_ok=True)

    actual_date = datetime.now().strftime("%Y-%m-%d")
    base_filename = f"reporte_{actual_date}"
    extension = ".docx"
    filename = base_filename + extension

    test = 1
    full_path = os.path.join(save_folder, filename)
    while os.path.exists(full_path):
        filename = f"{base_filename}_T{test}{extension}"
        full_path = os.path.join(save_folder, filename)
        test += 1

    # Crear documento
    doc = Document()
    doc.add_picture("FrED_Factory.png", width=Inches(7))

    p = doc.add_paragraph()
    run = p.add_run("Data Analysis of the FrED functioning")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(103, 24, 24)

    doc.add_paragraph("Hi! This was the report of today:")
    doc.add_paragraph(f"This test was made on {actual_date} and corresponds to test number {test} of the day.")
    doc.add_paragraph("Hope you have fun!")

    # Insertar imágenes del folder
    for imagen in os.listdir(image_folder):
        if imagen.lower().endswith((".png", ".jpg", ".jpeg")):
            ruta_imagen = os.path.join(image_folder, imagen)
            doc.add_paragraph(imagen)
            doc.add_picture(ruta_imagen, width=Inches(5.5))

    # Guardar documento en carpeta deseada
    doc.save(full_path)
    print(f"Document saved in {full_path}")

    # Eliminar imágenes después de insertarlas
    for imagen in os.listdir(image_folder):
        if imagen.lower().endswith((".png", ".jpg", ".jpeg")):
            os.remove(os.path.join(image_folder, imagen))
    print("Images eliminated.")

if __name__ == "__main__":
    image_folder = r"C:\Users\lench\Desktop\Data_Analyst\Images"
    create_docx(image_folder)
